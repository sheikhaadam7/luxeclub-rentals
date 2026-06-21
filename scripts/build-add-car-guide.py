"""
Generate the 'How to add a car to the website' Word doc.

One-off helper. Re-run after editing this script to refresh the file in
Downloads. Requires `python-docx` (pip install python-docx).
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUTPUT = Path.home() / "Downloads" / "LuxeClub - Add a Car to the Website.docx"

GOLD = RGBColor(0xB8, 0x8A, 0x3E)
GREY = RGBColor(0x55, 0x55, 0x55)
DARK = RGBColor(0x1A, 0x1A, 0x1A)

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

# Default body font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

def h1(text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.size = Pt(22)
    r.font.bold = True
    r.font.color.rgb = DARK

def h2(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(15)
    r.font.bold = True
    r.font.color.rgb = GOLD
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)

def h3(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.bold = True
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)

def para(*runs: tuple[str, dict]) -> None:
    """runs = list of (text, {'bold': True, 'mono': True, 'colour': RGBColor})"""
    p = doc.add_paragraph()
    for text, opts in runs:
        r = p.add_run(text)
        r.font.size = Pt(11)
        if opts.get("bold"):
            r.font.bold = True
        if opts.get("mono"):
            r.font.name = "Consolas"
            r.font.size = Pt(10)
        if opts.get("colour"):
            r.font.color.rgb = opts["colour"]

def text(s: str) -> None:
    para((s, {}))

def code_block(s: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(s)
    r.font.name = "Consolas"
    r.font.size = Pt(10)
    r.font.color.rgb = DARK
    p.paragraph_format.left_indent = Inches(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    # Light grey shading via XML
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="F4F4F4"/>')
    p._p.get_or_add_pPr().append(shd)

def bullet(text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        p.add_run(" " + text)
    else:
        p.add_run(text)

def note(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run("Note: ")
    r.font.bold = True
    r.font.color.rgb = GOLD
    p.add_run(text)
    p.paragraph_format.left_indent = Inches(0.2)


# ===== Title =====
h1("Add a Car to the Website")
p = doc.add_paragraph()
r = p.add_run("LuxeClub Rentals — Fleet Master Spreadsheet Workflow")
r.font.size = Pt(11)
r.font.color.rgb = GREY
r.font.italic = True

doc.add_paragraph()
text(
    "From start to finish: type two cells in Excel, drop your photos in a folder, "
    "run one command. The site catches up automatically at each step."
)


# ===== One-time setup =====
h2("One-time setup (skip if already done)")

h3("1.  Open a terminal in the project folder")
code_block("cd C:\\Users\\lenovo\\projects\\luxeclub-rentals")

h3("2.  Install the spreadsheet library")
code_block("pip install openpyxl")

h3("3.  Generate the master spreadsheet")
code_block("python scripts/export-fleet.py")
text("This creates data/fleet.xlsx with one row per car already in the database.")


# ===== Daily workflow =====
h2("Every time you want to add a car")

h3("Step 1  —  Start the watcher (leave running in its own terminal)")
code_block("python scripts/watch-fleet.py")
text(
    "The watcher detects every save of data/fleet.xlsx and automatically pushes "
    "the changes to the live site. Leave this window open the whole time you're "
    "editing. Press Ctrl+C to stop it when you're done."
)
note(
    "If you'd rather not run the watcher, you can manually run "
    "'python scripts/import-fleet.py --apply' after each save instead."
)


h3("Step 2  —  Add a row to data/fleet.xlsx")
text("Open data/fleet.xlsx in Excel. Scroll to the bottom and add a new row.")
para(("Minimum required cells: ", {"bold": True}))
bullet("Lamborghini Aventador SVJ  (or whatever the car is called)", bold_prefix="name")
bullet("3000  (in AED per day)", bold_prefix="daily_rate")
text("Leave the slug cell blank — it will auto-derive from the name.")

para(("Optional cells worth filling in: ", {"bold": True}))
bullet("Lamborghini, Bentley, Audi… (drives the brand filter on the catalogue page)", bold_prefix="brand")
bullet("2024", bold_prefix="year")
bullet("weekly_rate / monthly_rate", bold_prefix="prices")
bullet("AED amount", bold_prefix="deposit_amount")
bullet("AED charged per km when the renter exceeds the included mileage", bold_prefix="overage_rate_per_km")
bullet("engine / horsepower / seats / transmission / drivetrain", bold_prefix="specs")
bullet("X in the matching columns (Sports / SUV / Convertible / Sedan / Coupe / Family)", bold_prefix="categories")


h3("Step 3  —  Save the spreadsheet (Ctrl + S)")
text("The watcher fires the import automatically. You'll see something like this:")
code_block(
    "+ lamborghini-aventador-svj: NEW (name='Lamborghini Aventador SVJ', daily_rate=3000)\n"
    "  [no primary_image_url → will be hidden until photos uploaded]\n"
    "OK   + lamborghini-aventador-svj\n"
    "     Folder ready for photos: data/images-to-upload/2026-06-21_lamborghini-aventador-svj/"
)
para(("What just happened: ", {"bold": True}))
bullet("Car was inserted into the database.")
bullet("A dated folder was auto-created under data/images-to-upload/ with a README.txt inside.")
bullet("The car is HIDDEN from the public site until photos are uploaded. This is on purpose — no placeholder cards on the public catalogue.")


h3("Step 4  —  Drop the photos into the auto-created folder")
text(
    "Open File Explorer and navigate to the folder that was just created — "
    "e.g. data/images-to-upload/2026-06-21_lamborghini-aventador-svj/"
)
para(("Photo naming rules: ", {"bold": True}))
bullet("Must be named primary.jpg, primary.png, or primary.webp. This is the hero image shown first everywhere.", bold_prefix="The main / hero photo")
bullet("Anything else (1.jpg, 2.jpg, gallery-a.png …). They become the gallery in alphabetical order.", bold_prefix="Other photos")


h3("Step 5  —  Upload the photos")
code_block("python scripts/upload-fleet-images.py --apply")
para(("What happens: ", {"bold": True}))
bullet("Photos are uploaded to Supabase Storage.")
bullet("The database row gets primary_image_url and image_urls filled in.")
bullet("The car is auto-published — it appears on the public catalogue immediately.")
bullet("Previous URLs (if you're replacing existing photos) are backed up to data/images-backup-*.csv just in case.")


h3("Step 6 (optional)  —  Auto-fill the description with Claude")
code_block("python scripts/enhance-fleet.py --apply")
text(
    "Finds every car with a blank description (including the new one) and asks "
    "Claude for a 2–3 sentence rental write-up. The script automatically rejects "
    "anything that mentions 24/7, hard driving, unlimited mileage, or other "
    "phrases you've said to avoid. Clean descriptions go straight to the DB; "
    "flagged ones land in data/descriptions-for-review.txt for you to review."
)


# ===== Quick reference =====
h2("Quick reference — common tasks")

h3("Change a price")
text("Edit the cell in daily_rate / weekly_rate / monthly_rate. Save. Done.")

h3("Take a car off the site temporarily (e.g. sold out for a week)")
text("Clear the X in the bookable? column. Save. The car stays visible but shows the 'temporarily unavailable' message.")
text("Put it back: type X in bookable?. Save.")

h3("Retire a car permanently (or hide it from the public site)")
text("Type X in the retire? column. Save. The car disappears from the catalogue immediately. Booking history is preserved.")
text("Un-retire: clear the X. Save.")

h3("Replace a car's photos")
text(
    "Create the folder data/images-to-upload/{slug}/ (no date prefix needed), "
    "drop the new photos in (rename hero to primary.jpg), and run:"
)
code_block("python scripts/upload-fleet-images.py --apply")
text("Old photos are automatically deleted from storage. URLs are backed up first.")


# ===== Row colours =====
h2("What the row colours mean")
bullet("Car is on the public site and bookable.", bold_prefix="No colour (Live)")
bullet("Car is in the database but has no hero photo yet. Hidden from the public site.", bold_prefix="Yellow (Awaiting photos)")
bullet("Car is visible on the site with the 'temporarily unavailable' message. Bookable? is empty.", bold_prefix="Light orange (Sold out)")
bullet("Hidden from the public site. Retire? has an X.", bold_prefix="Grey (Retired)")


# ===== Troubleshooting =====
h2("If something goes wrong")

h3("\"slug ... not in vehicles\" when uploading images")
text(
    "The folder name doesn't match a row in the database. Did you add the row "
    "to the spreadsheet and save first? You have to do that before dropping the "
    "photos in."
)

h3("\"BULK RETIRE WARNING\" when saving the spreadsheet")
text(
    "You accidentally marked X in retire? on too many cars. The script refuses "
    "to retire more than 30% of the fleet in one go without --force. Check what "
    "happened, undo if needed, and save again."
)

h3("The watcher isn't picking up my save")
text(
    "Look at the watcher terminal. If it says 'mtime changed but content "
    "identical', it means you saved without actually changing anything — that's "
    "fine. Make an actual edit and save again."
)

h3("Don't run the legacy scraper")
text(
    "scripts/scraper/scrape-vehicles.ts will overwrite your spreadsheet edits. "
    "Don't run 'npm run scrape' while you're using the spreadsheet workflow."
)


# ===== Outro =====
doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run("That's it. Two cells in Excel + one upload command = car on the website.")
r.font.italic = True
r.font.color.rgb = GREY


OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(f"Wrote {OUTPUT}")
